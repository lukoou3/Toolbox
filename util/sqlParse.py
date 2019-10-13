import re
import os

from bs4 import BeautifulSoup
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from util import util

class Field:
    def __init__(self, name='', type='', dtype='', comment=''):
        self.name = name
        self.type = type
        self.dtype = dtype
        self.comment = comment
        self.is_not_null = False
        self.is_key = False
        self.is_auto_increase = False

    def __repr__(self):
        return "Field(name=%s, type=%s, dtype=%s, comment=%s, is_not_null=%s, is_key=%s, is_auto_increase=%s)" % \
               (self.name, self.type, self.dtype, self.comment, self.is_not_null, self.is_key, self.is_auto_increase)

class Table:
    def __init__(self):
        self.name = ''
        self.comment = ''
        self.fields = []

    def __repr__(self):
        return "Table(name=%s, comment=%s, fields=%s)" % (self.name, self.comment, repr(self.fields))

    def pprint(self):
        import pprint
        print("Table(\nname=%s,"% self.name)
        print("comment=%s," % self.comment)
        print("fields=")
        pprint.pprint(self.fields)
        print(")")

def parse_sql_table(sql_text):
    """
    sql文本转Table对象列表
    :param sql_text: sql文本
    :return: list of Table
    """
    tables = []
    ret = re.findall(r"""create\s+(?:external\s+|)table\s+
                    `{0,1}(.+?)`{0,1}\s*    #表名
                    \(([^;]+)\)           #表内容
                    ([^;]*?)          #表描述
                    ;""", sql_text, re.X | re.I)
    table_comm_re = re.compile(r"COMMENT='(.*?)'", re.I)
    for per_table_ret in ret:
        table_name = per_table_ret[0]
        table_body = per_table_ret[1]
        table_desc = per_table_ret[2]
        table_body_lines = list(map(lambda x: x.strip(), table_body.strip().splitlines()))
        new_table = Table()
        # 遍历( ... )里面的每一行
        key_names = []
        for line in table_body_lines:
            line_lower = line.lower()
            if 'primary key' in line_lower:
                key_names.append(line[line.find('(') + 1:line.find(')')])
            if 'key ' not in line_lower:
                ret_line = re.search(r'`{0,1}([\w\d_]+)`{0,1}\s+([^\s,]+)\s*(.*)', line)
                field_name = ret_line.group(1)
                field_type = ret_line.group(2)
                field_dtype = field_type.split("(")[0].strip()
                field_comment = ''
                if 'comment ' in line_lower:
                    field_comment = re.search(r"comment\s+'(.*?)'", line, re.I).group(1)
                new_field = Field(field_name, field_type, field_dtype, field_comment)
                if 'not null' in line_lower:
                    new_field.is_not_null = True
                if 'auto_increment' in line_lower:
                    new_field.is_auto_increase = True
                new_table.fields.append(new_field)
        # 处理primary key
        for per_field in new_table.fields:
            if per_field.name in key_names:
                per_field.is_key = True
        table_comment = ''
        ret_comment = table_comm_re.search(table_desc)
        if ret_comment is not None:
            table_comment = ret_comment.group(1)
        new_table.comment = table_comment
        new_table.name = table_name
        tables.append(new_table)
    return tables

def sql_to_javabean(sql_text, useWrapperClass=False):
    """
    sql文本转javabean，
    :param sql_text: sql文本
    :param useWrapperClass: 属性是否使用包装类
    :return: list of (class name, javaBean text)
    """
    tables = parse_sql_table(sql_text)
    return [table_to_javabean(table, useWrapperClass) for table in tables]

def sql_to_scalacase(sql_text):
    tables = parse_sql_table(sql_text)
    return [table_to_scalacase(table) for table in tables]

def sql_to_scalabean(sql_text):
    tables = parse_sql_table(sql_text)
    return [table_to_scalabean(table) for table in tables]

def sql_to_docx(sql_text,path="sql.docx",explorer=False):
    tables = parse_sql_table(sql_text)

    NORMAL_FONT = '微软雅黑'
    NORMAL_FONT_SIZE = 10
    TABLE_STYLE = 'Light Grid Accent 3'

    document = Document()
    # 样式设置
    document.styles['Normal'].font.name = NORMAL_FONT
    document.styles['Normal'].font.size = Pt(NORMAL_FONT_SIZE)
    # noinspection PyProtectedMember
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), NORMAL_FONT)
    # 1.Head生成
    doc_title = "数据库表"
    head_runs = document.add_heading(doc_title, 0).runs
    for per_run in head_runs:
        per_run.font.name = NORMAL_FONT
        # noinspection PyProtectedMember
        per_run._element.rPr.rFonts.set(qn('w:eastAsia'), NORMAL_FONT)
    # 2.Sql表格的生成
    for sql_table in tables:
        # 表Title生成
        tag_text = sql_table.name + '(' + sql_table.comment + ')' if not sql_table.comment == 'null' else sql_table.name
        label = document.add_paragraph()
        label_table_name = label.add_run(tag_text[0].upper() + tag_text[1:])
        label_table_name.font.size = Pt(14)
        # 表格生成
        row, col = len(sql_table.fields) + 1, 6
        docx_table = document.add_table(rows=row, cols=col, style=TABLE_STYLE)
        col_index = 0
        docx_table.cell(0, col_index).text = '字段名称'
        col_index += 1
        docx_table.cell(0, col_index).text = '数据类型'
        col_index += 1
        docx_table.cell(0, col_index).text = '主键'
        col_index += 1
        docx_table.cell(0, col_index).text = '不可空'
        col_index += 1
        docx_table.cell(0, col_index).text = '自增'
        col_index += 1
        docx_table.cell(0, col_index).text = '字段注释'
        for field, row_index in zip(sql_table.fields, range(len(sql_table.fields))):
            col_index = 0
            docx_table.cell(row_index + 1, col_index).text = field.name
            col_index += 1
            docx_table.cell(row_index + 1, col_index).text = field.type
            col_index += 1
            docx_table.cell(row_index + 1, col_index).text = str(field.is_key)
            col_index += 1
            docx_table.cell(row_index + 1, col_index).text = str(field.is_not_null)
            col_index += 1
            docx_table.cell(row_index + 1, col_index).text = str(field.is_auto_increase)
            col_index += 1
            docx_table.cell(row_index + 1, col_index).text = field.comment
        document.add_paragraph("\n")
    document.save(path)
    if explorer:
        os.system("explorer %s" % (path))

def sql_to_html(sql_text,path="sql.html",explorer=False):
    html_css = """
        th {
            background-color: rgb(81, 130, 187);
            color: #fff;
            border-bottom-width: 0;
        }
        td {
            color: #000;
        }
        tr, th {
            border-width: 1px;
            border-style: solid;
            border-color: rgb(81, 130, 187);
        }
        td, th {
            padding: 5px 10px;
            font-size: 12px;
            font-family: Verdana;
            font-weight: bold;
        }
        table {
            border-width: 1px;
            border-collapse: collapse;
            float: left;
            margin: 10px;
        }
    """
    html = """
    <html><head><title>PdmShow</title><style>%s</style></head><body></body></html>
    """ % html_css
    tables = parse_sql_table(sql_text)

    html = BeautifulSoup(html, 'lxml')
    # 遍历数据表
    for table in tables:
        table1 = html.new_tag(name='table')
        # 标题行
        tr_head = html.new_tag(name='tr')
        td_head = html.new_tag(name='th', colspan="3")
        td_head.append(table.name + '(' + table.comment + ')')
        tr_head.append(td_head)
        table1.append(tr_head)
        # field行
        for field in table.fields:
            tr_field = html.new_tag(name='tr')
            td_name = html.new_tag(name='td')
            td_name.append(field.name)
            td_type = html.new_tag(name='td')
            td_type.append(field.type)
            td_comment = html.new_tag(name='td')
            td_comment.append(field.comment)
            tr_field.append(td_name)
            tr_field.append(td_type)
            tr_field.append(td_comment)
            table1.append(tr_field)
        html.body.append(table1)

    result_html = html.prettify()
    with open(path, "w", encoding="utf-8") as fp:
        fp.write(result_html)
    if explorer:
        os.system("explorer %s" % (path))

def table_to_javabean(table, useWrapperClass=False):
    """
    Table对象转为javaBean
    :param table: Table obj
    :return: (class name, javaBean text)
    """
    type_map = {"varchar": "String", "datetime": "Date", "bigint": "long", "smallint": "int",
                "tinyint": "int", "int": "int", "float": "float", "double": "double"}
    wrapper_class_map = {"int": "Integer", "long": "Long", "double": "Double"}
    class_fmt = """public class {class_name} {{
{field_content}
}}
    """
    field_fmt = """    private {type} {name};{comment}"""
    if useWrapperClass:
        type_map.update(wrapper_class_map)
    field_list = [(type_map.get(field.dtype, "String"), util.fieldname_under2camel(field.name), field.comment) for field in table.fields]
    class_name = util.fieldname_under2camel(table.name, True)
    field_content = "\n".join( [field_fmt.format(type=type,name=name,comment="//"+comment if comment else comment) for type,name,comment in field_list] )
    class_text = class_fmt.format(**{"class_name": class_name,
                                   "field_content": field_content })
    return class_name, class_text

def table_to_scalacase(table):
    type_map = {"varchar": "String", "datetime": "Date", "bigint": "Long", "smallint": "Int",
                "tinyint": "Int", "int": "Int", "float": "Float", "double": "Double"}
    class_fmt = """case class {}({})"""
    class_name = util.fieldname_under2camel(table.name, True)
    field_list = [(type_map.get(field.dtype, "String"), util.fieldname_under2camel(field.name), field.comment) for field in table.fields]
    class_text = class_fmt.format(class_name, ", ".join( ["{}:{}".format(name,type) for type,name,comment in field_list] ))
    return class_name, class_text

def table_to_scalabean(table):
    type_map = {"varchar": "String", "datetime": "Date", "bigint": "Long", "smallint": "Int",
                "tinyint": "Int", "int": "Int", "float": "Float", "double": "Double"}
    class_fmt = """class {class_name}{{
{field_content}
}}"""
    field_fmt = """  var {name}: {type} = _{comment}"""
    class_name = util.fieldname_under2camel(table.name, True)
    field_list = [(type_map.get(field.dtype, "String"), util.fieldname_under2camel(field.name), field.comment) for field
                  in table.fields]
    field_content = "\n".join([field_fmt.format(type=type, name=name, comment="//" + comment if comment else comment)
         for type, name, comment in field_list])
    class_text = class_fmt.format(class_name=class_name,field_content=field_content)
    return class_name, class_text

def sql_to_sqlalchemy(sql_text):
    table_texts = re.findall(r"""create\s+(?:external\s+|)table\s+
                        `{0,1}.+?`{0,1}\s*    #表名
                        \([^;]+\)           #表内容
                        [^;]*?              #表描述
                        ;""", sql_text, re.X | re.I)

    return [sql_table_to_sqlalchemy(table_text) for table_text in table_texts]

def sql_table_to_sqlalchemy(text):
    type_map = {"varchar": "String", "datetime": "DateTime", "bigint": "BigInteger", "smallint": "SmallInteger",
                "tinyint": "SmallInteger", "text": "Text", "int": "Integer", "double": "Float", "char": "String",
                "set": "Enum"}
    s = []
    primary_key_l = []
    lines = text.split("\n")  # 表设计行拆分
    for line in lines[::-1]:  # 遍历表设计行
        j = line.strip().split(" ")  # 倒序遍历，并按空格切分
        if len(j) > 2:  # 只关注行长度超过2的元素
            column = j[0].replace("`", "")
            i_type = j[1]
            if column == "PRIMARY":
                primary_key_l = re.sub(r'`|\(|\)', '', j[2]).split(",")  # 拿到主键key
                continue
            elif column == "CREATE":  # 获取表名
                table_name = j[2].replace("`", "")
                s.append("    " + '__tablename__ = "%s"' % table_name)
                s.append("class %s(Base):" % table_name)
                continue
            elif column in ("UNIQUE", ")", "KEY"):  # 非表列名，跳过
                continue
            if i_type in type_map.keys():  # 类型存在映射表中
                i_type = i_type.replace(i_type, type_map[i_type]) + "()"
            elif "(" in i_type and i_type.split("(")[0] in type_map.keys():  # 类型有长度声明，提取类型字段，找到映射表映射value，并替换
                old_type = i_type.split("(")[0]
                new_type = type_map[old_type]
                i_type = i_type.replace(old_type, new_type)
            else:
                i_type = i_type.replace(i_type, type_map[i_type]) + "()"
                print("Catch any case not in type_map:%s" % i_type)

            if column in primary_key_l:  # 列名存在主键数组中
                i_type = i_type + ", primary_key=True"
            s.append("    " + column + " = Column(" + i_type + ")")

    return table_name,"\n".join(s[::-1])  # 反序输出

if __name__ == '__main__':
    sql_text = """CREATE TABLE `test_rerm` (
  `id` int(10) unsigned NOT NULL AUTO_INCREMENT,
  `ip` int(10) unsigned NOT NULL,
  `termType` varchar(32) NOT NULL DEFAULT '' comment '1111',
  `vendor` varchar(32) NOT NULL,
  `date` timestamp NULL DEFAULT NULL ON UPDATE CURRENT_TIMESTAMP,
  PRIMARY KEY (`id`)
) ENGINE=InnoDB AUTO_INCREMENT=3001 DEFAULT CHARSET=utf8 COMMENT='终端类型表';"""
    sql_to_html(sql_text)
    for name,text in sql_to_scalabean(sql_text):
        print(text)
